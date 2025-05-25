import boto3
import json
import base64
import os
import sys
import argparse
import time
from openpyxl.utils import get_column_letter
from botocore.exceptions import ClientError

# Try to import requests, but don't fail if it's not available
try:
    import requests
except ImportError:
    requests = None

def get_media_type(file_path):
    """Determine the media type based on file extension"""
    _, file_extension = os.path.splitext(file_path)
    file_extension = file_extension.lower()
    
    if file_extension in ['.jpg', '.jpeg', '.png']:
        return "image/jpeg" if file_extension in ['.jpg', '.jpeg'] else "image/png"
    elif file_extension == '.pdf':
        return "application/pdf"
    else:
        print(f"Unsupported file extension: {file_extension}")
        print("Supported formats: .jpg, .jpeg, .png, .pdf")
        sys.exit(1)

def check_and_install_fonts():
    """Check for and install required Arabic/Urdu fonts"""
    import tempfile
    import subprocess
    import requests
    from pathlib import Path
    
    # Define font paths
    user_fonts_dir = os.path.expanduser("~/Library/Fonts")
    
    # Create fonts directory if it doesn't exist
    os.makedirs(user_fonts_dir, exist_ok=True)
    
    # Define required fonts and their direct download URLs (no zip files)
    required_fonts = {
        "Amiri-Regular.ttf": "https://github.com/alif-type/amiri/raw/master/amiri-regular.ttf",
        "NotoSansArabic-Regular.ttf": "https://github.com/googlefonts/noto-fonts/raw/main/hinted/ttf/NotoSansArabic/NotoSansArabic-Regular.ttf",
        "NotoNastaliqUrdu-Regular.ttf": "https://github.com/googlefonts/noto-fonts/raw/main/hinted/ttf/NotoNastaliqUrdu/NotoNastaliqUrdu-Regular.ttf"
    }
    
    # Check which fonts are missing
    missing_fonts = {}
    for font_file, url in required_fonts.items():
        if not os.path.exists(os.path.join(user_fonts_dir, font_file)):
            missing_fonts[font_file] = url
    
    if not missing_fonts:
        print("All required Arabic and Urdu fonts are already installed.")
        return
    
    print(f"Installing {len(missing_fonts)} missing Arabic/Urdu fonts...")
    
    # Download missing fonts directly
    for font_file, url in missing_fonts.items():
        font_path = os.path.join(user_fonts_dir, font_file)
        print(f"Downloading {font_file} from {url}...")
        
        try:
            # Method 1: Using requests (more reliable)
            try:
                response = requests.get(url, timeout=30)
                response.raise_for_status()  # Raise exception for 4XX/5XX responses
                
                with open(font_path, 'wb') as f:
                    f.write(response.content)
                    
                print(f"Successfully installed {font_file}")
                continue  # Skip to next font if successful
            except (requests.RequestException, ImportError) as e:
                print(f"Requests download failed for {font_file}, trying curl: {e}")
            
            # Method 2: Using curl as fallback
            subprocess.run(
                ["curl", "-L", url, "-o", font_path, "--connect-timeout", "10", "--max-time", "30"],
                check=True, capture_output=True, text=True, timeout=45
            )
            print(f"Successfully installed {font_file} using curl")
            
        except subprocess.CalledProcessError as e:
            print(f"Error installing {font_file}: {e}")
            print(f"stdout: {e.stdout}")
            print(f"stderr: {e.stderr}")
        except subprocess.TimeoutExpired:
            print(f"Timeout while downloading {font_file}")
        except Exception as e:
            print(f"Unexpected error installing {font_file}: {e}")
    
    # Verify installation
    installed_count = 0
    for font_file in required_fonts:
        if os.path.exists(os.path.join(user_fonts_dir, font_file)):
            installed_count += 1
    
    if installed_count == len(required_fonts):
        print("All required Arabic and Urdu fonts have been successfully installed.")
    else:
        print(f"Installed {installed_count}/{len(required_fonts)} required fonts.")
        print("Some fonts may need to be installed manually.")
        
    # On macOS, clear font cache
    try:
        subprocess.run(["atsutil", "databases", "-removeUser"], 
                      check=False, capture_output=True, timeout=10)
    except Exception as e:
        print(f"Note: Could not clear font cache: {e}")
        pass

def parse_arguments():
    """Parse command line arguments"""
    parser = argparse.ArgumentParser(description='Extract text from images or PDFs using Amazon Bedrock')
    parser.add_argument('file_path', nargs='?', default='/Users/test/Downloads/UAE-Invoice-Template-1.jpg',
                        help='Path to the image or PDF file')
    parser.add_argument('--region', default='us-west-2', help='AWS region for Bedrock')
    parser.add_argument('--profile', default='default', help='AWS profile name')
    parser.add_argument('--model', default='anthropic.claude-3-5-sonnet-20240620-v1:0', 
                        help='Bedrock model ID to use')
    parser.add_argument('--formats', default='txt,pdf,docx,xlsx', 
                        help='Output formats (comma-separated: txt,pdf,docx,xlsx)')
    return parser.parse_args()

def save_as_text(text, output_path):
    """Save extracted text to a text file"""
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(text)
    print(f"Text saved to: {output_path}")

def save_as_pdf(text, output_path):
    """Placeholder function for PDF generation - disabled due to compatibility issues"""
    print(f"PDF generation has been disabled due to compatibility issues.")
    print(f"To generate PDFs, please use an alternative tool or library.")

def save_as_docx(text, output_path):
    """Save extracted text to a Word document with special formatting for financial documents"""
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import re
        
        # Create document
        doc = Document()
        
        # Detect if this is a financial document
        financial_keywords = ['invoice', 'bill', 'receipt', 'statement', 'purchase order', 'po', 'ledger', 
                             'فاتورة', 'حساب', 'إيصال', 'أمر شراء', 'دفتر الأستاذ']
        
        is_financial = any(keyword in text.lower() for keyword in financial_keywords)
        
        if is_financial:
            print("Financial document detected. Applying specialized formatting to Word document...")
            
            # Process text into structured data
            lines = text.split('\n')
            
            # Skip the "Here is the text extracted from the document:" line if present
            if lines and lines[0].startswith("Here is the text extracted from"):
                lines = lines[1:]
            
            # Extract header information, items, and totals
            header_section = []
            item_section = []
            total_section = []
            
            # Simple state machine to categorize lines
            in_items = False
            in_totals = False
            
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                
                # Look for item section indicators
                if re.search(r'(البيان|item|description|qty|quantity|سعر|price|amount|المبلغ)', line.lower()):
                    in_items = True
                    in_totals = False
                    item_section.append(line)
                # Look for total section indicators
                elif re.search(r'(total|subtotal|المجموع|الإجمالي|vat|ضريبة|tax)', line.lower()) and in_items:
                    in_items = False
                    in_totals = True
                    total_section.append(line)
                # Categorize based on current state
                elif in_items:
                    item_section.append(line)
                elif in_totals:
                    total_section.append(line)
                else:
                    header_section.append(line)
            
            # Add title
            title = doc.add_heading('', level=1)
            title_run = title.add_run("Invoice / Financial Document")
            title_run.bold = True
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add header section
            doc.add_heading('Header Information', level=2)
            table = doc.add_table(rows=0, cols=2)
            table.style = 'Table Grid'
            
            for line in header_section:
                if ':' in line:
                    key, value = line.split(':', 1)
                    row_cells = table.add_row().cells
                    row_cells[0].text = key.strip()
                    row_cells[1].text = value.strip()
                else:
                    # For lines without a colon, span across both columns
                    row_cells = table.add_row().cells
                    row_cells[0].merge(row_cells[1])
                    row_cells[0].text = line
            
            # Add item section if we have items
            if item_section:
                doc.add_heading('Items', level=2)
                
                # Try to determine the number of columns from the first line
                first_line = item_section[0] if item_section else ""
                num_columns = 1
                
                # Check if the first line might contain column headers
                if re.search(r'(description|qty|price|amount|البيان|الكمية|السعر|المبلغ)', first_line.lower()):
                    # Count potential column headers
                    column_indicators = ['item', 'description', 'qty', 'quantity', 'price', 'amount', 'total',
                                        'البيان', 'الوصف', 'الكمية', 'السعر', 'المبلغ']
                    count = sum(1 for indicator in column_indicators if indicator.lower() in first_line.lower())
                    num_columns = max(count, 2)  # At least 2 columns
                
                # Create items table
                items_table = doc.add_table(rows=0, cols=num_columns)
                items_table.style = 'Table Grid'
                
                # Add header row if we detected column headers
                if num_columns > 1:
                    header_row = items_table.add_row().cells
                    
                    # Try to extract column headers
                    headers = []
                    for indicator in ['description', 'qty', 'price', 'amount']:
                        if indicator.lower() in first_line.lower():
                            headers.append(indicator.capitalize())
                    
                    # If we couldn't extract enough headers, use defaults
                    while len(headers) < num_columns:
                        headers.append(f"Column {len(headers)+1}")
                    
                    for i, header in enumerate(headers):
                        header_row[i].text = header
                        for paragraph in header_row[i].paragraphs:
                            for run in paragraph.runs:
                                run.bold = True
                
                    # Skip the header line in the data
                    item_data = item_section[1:]
                else:
                    item_data = item_section
                
                # Add item rows
                for line in item_data:
                    if not line.strip():
                        continue
                        
                    row_cells = items_table.add_row().cells
                    
                    if num_columns > 1:
                        # Try to split the line into columns
                        parts = []
                        
                        # Check if the line has numbers that might indicate price/quantity
                        numbers = re.findall(r'\d+(?:,\d+)*(?:\.\d+)?', line)
                        
                        if numbers and len(numbers) >= 2:
                            # This might be a line with quantity and price
                            # Extract the description (text before the first number)
                            desc_match = re.split(r'\d+(?:,\d+)*(?:\.\d+)?', line)[0].strip()
                            if desc_match:
                                parts.append(desc_match)
                                
                            # Add the numbers as quantity, price, etc.
                            for num in numbers:
                                parts.append(num)
                        else:
                            # Just add the whole line as description
                            parts.append(line)
                            
                        # Fill the cells
                        for i, part in enumerate(parts):
                            if i < num_columns:
                                row_cells[i].text = part
                    else:
                        # Single column - just add the whole line
                        row_cells[0].text = line
            
            # Add total section
            if total_section:
                doc.add_heading('Totals', level=2)
                totals_table = doc.add_table(rows=0, cols=2)
                totals_table.style = 'Table Grid'
                
                for line in total_section:
                    if ':' in line:
                        key, value = line.split(':', 1)
                        row_cells = totals_table.add_row().cells
                        row_cells[0].text = key.strip()
                        row_cells[1].text = value.strip()
                    else:
                        # Check if this line contains a currency amount
                        amount_match = re.search(r'([\d,\.]+)\s*([A-Z]{3})', line)
                        if amount_match:
                            row_cells = totals_table.add_row().cells
                            row_cells[0].text = line.replace(amount_match.group(0), '').strip()
                            row_cells[1].text = amount_match.group(0)
                        else:
                            row_cells = totals_table.add_row().cells
                            row_cells[0].merge(row_cells[1])
                            row_cells[0].text = line
                            
                            # Make the text bold if it looks like a grand total
                            if re.search(r'(total|المجموع|الإجمالي)', line.lower()):
                                for paragraph in row_cells[0].paragraphs:
                                    for run in paragraph.runs:
                                        run.bold = True
            
            # Add footer if there's any text after the totals
            remaining_text = []
            capture_remaining = False
            for line in lines:
                if capture_remaining:
                    remaining_text.append(line)
                elif line.strip() in [l.strip() for l in total_section]:
                    capture_remaining = True
            
            if remaining_text and any(line.strip() for line in remaining_text):
                doc.add_heading('Additional Information', level=2)
                for line in remaining_text:
                    if line.strip():
                        doc.add_paragraph(line)
        
        else:
            # Standard processing for non-financial documents
            # Add text paragraphs
            for line in text.split('\n'):
                if line.strip():  # Skip empty lines
                    doc.add_paragraph(line)
                else:
                    doc.add_paragraph()  # Add empty paragraph for blank lines
        
        # Save the document
        doc.save(output_path)
        print(f"Word document saved to: {output_path}")
        
    except ImportError:
        print("Warning: python-docx package not installed. Word output skipped.")
        print("To install, run: pip install python-docx")

def save_as_excel(text, output_path):
    """Save extracted text to an Excel file with special formatting for financial documents"""
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, Alignment, Border, Side
        import re
        
        # Create workbook and select active sheet
        wb = Workbook()
        ws = wb.active
        ws.title = "Extracted Text"
        
        # Detect if this is a financial document
        financial_keywords = ['invoice', 'bill', 'receipt', 'statement', 'purchase order', 'po', 'ledger', 
                             'فاتورة', 'حساب', 'إيصال', 'أمر شراء', 'دفتر الأستاذ']
        
        is_financial = any(keyword in text.lower() for keyword in financial_keywords)
        
        if is_financial:
            print("Financial document detected. Applying specialized formatting...")
            
            # Define styles
            header_font = Font(bold=True)
            centered_alignment = Alignment(horizontal='center')
            right_alignment = Alignment(horizontal='right')
            border = Border(
                left=Side(style='thin'),
                right=Side(style='thin'),
                top=Side(style='thin'),
                bottom=Side(style='thin')
            )
            
            # Process text into structured data
            lines = text.split('\n')
            row_idx = 1
            
            # Skip the "Here is the text extracted from the document:" line if present
            if lines and lines[0].startswith("Here is the text extracted from"):
                lines = lines[1:]
            
            # Extract header information (date, invoice number, etc.)
            header_section = []
            item_section = []
            total_section = []
            
            # Simple state machine to categorize lines
            in_items = False
            in_totals = False
            
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                
                # Look for item section indicators
                if re.search(r'(البيان|item|description|qty|quantity|سعر|price|amount|المبلغ)', line.lower()):
                    in_items = True
                    in_totals = False
                    item_section.append(line)
                # Look for total section indicators
                elif re.search(r'(total|subtotal|المجموع|الإجمالي|vat|ضريبة|tax)', line.lower()) and in_items:
                    in_items = False
                    in_totals = True
                    total_section.append(line)
                # Categorize based on current state
                elif in_items:
                    item_section.append(line)
                elif in_totals:
                    total_section.append(line)
                else:
                    header_section.append(line)
            
            # Write header section
            for i, line in enumerate(header_section, 1):
                if ':' in line:
                    key, value = line.split(':', 1)
                    ws.cell(row=i, column=1, value=key.strip())
                    ws.cell(row=i, column=2, value=value.strip())
                else:
                    ws.cell(row=i, column=1, value=line)
                    ws.merge_cells(f'A{i}:D{i}')
            
            row_idx = len(header_section) + 2  # Skip a row
            
            # Write item header if we detected items
            if item_section:
                # Try to detect columns from the first line
                columns = []
                first_item_line = item_section[0]
                
                # Check for common column headers
                potential_columns = ['item', 'description', 'qty', 'quantity', 'price', 'amount', 'total',
                                    'البيان', 'الوصف', 'الكمية', 'السعر', 'المبلغ']
                
                for col in potential_columns:
                    if col.lower() in first_item_line.lower():
                        columns.append(col)
                
                # If we couldn't detect columns, use default ones
                if not columns:
                    columns = ['Description', 'Quantity', 'Price', 'Amount']
                
                # Write column headers
                for i, col in enumerate(columns, 1):
                    cell = ws.cell(row=row_idx, column=i, value=col)
                    cell.font = header_font
                    cell.alignment = centered_alignment
                    cell.border = border
                
                row_idx += 1
                
                # Write item rows
                for line in item_section[1:]:  # Skip the header line
                    # Try to split the line into columns
                    parts = []
                    
                    # Check if the line has numbers that might indicate price/quantity
                    numbers = re.findall(r'\d+(?:,\d+)*(?:\.\d+)?', line)
                    
                    if numbers and len(numbers) >= 2:
                        # This might be a line with quantity and price
                        # Extract the description (text before the first number)
                        desc_match = re.split(r'\d+(?:,\d+)*(?:\.\d+)?', line)[0].strip()
                        if desc_match:
                            parts.append(desc_match)
                            
                        # Add the numbers as quantity, price, etc.
                        for num in numbers:
                            parts.append(num)
                    else:
                        # Just add the whole line as description
                        parts.append(line)
                    
                    # Write the parts to columns
                    for i, part in enumerate(parts, 1):
                        cell = ws.cell(row=row_idx, column=i, value=part)
                        cell.border = border
                        
                        # Right-align if it looks like a number
                        if re.match(r'^[\d,\.]+$', part.strip()):
                            cell.alignment = right_alignment
                    
                    row_idx += 1
            
            # Write total section
            row_idx += 1  # Skip a row
            for line in total_section:
                if ':' in line:
                    key, value = line.split(':', 1)
                    cell1 = ws.cell(row=row_idx, column=1, value=key.strip())
                    cell2 = ws.cell(row=row_idx, column=2, value=value.strip())
                    cell1.font = header_font
                    cell2.alignment = right_alignment
                else:
                    # Check if this line contains a currency amount
                    amount_match = re.search(r'([\d,\.]+)\s*([A-Z]{3})', line)
                    if amount_match:
                        amount, currency = amount_match.groups()
                        ws.cell(row=row_idx, column=1, value=line.replace(amount_match.group(0), '').strip())
                        cell = ws.cell(row=row_idx, column=2, value=f"{amount} {currency}")
                        cell.alignment = right_alignment
                    else:
                        ws.cell(row=row_idx, column=1, value=line)
                
                row_idx += 1
            
            # Adjust column widths
            for col_idx in range(1, ws.max_column + 1):
                max_length = 0
                for row_idx in range(1, ws.max_row + 1):
                    cell = ws.cell(row=row_idx, column=col_idx)
                    if cell.value:
                        cell_length = len(str(cell.value))
                        if cell_length > max_length:
                            max_length = cell_length
                adjusted_width = (max_length + 2)
                ws.column_dimensions[get_column_letter(col_idx)].width = adjusted_width
        
        else:
            # Standard processing for non-financial documents
            lines = text.split('\n')
            for row_idx, line in enumerate(lines, 1):
                # Try to split by common delimiters
                if ':' in line:
                    # This might be a key-value pair
                    parts = line.split(':', 1)
                    ws.cell(row=row_idx, column=1, value=parts[0].strip())
                    if len(parts) > 1:
                        ws.cell(row=row_idx, column=2, value=parts[1].strip())
                else:
                    # Just put the whole line in column A
                    ws.cell(row=row_idx, column=1, value=line)
        
        # Save the workbook
        wb.save(output_path)
        print(f"Excel file saved to: {output_path}")
        
    except ImportError:
        print("Warning: openpyxl package not installed. Excel output skipped.")
        print("To install, run: pip install openpyxl")
    except Exception as e:
        print(f"Unexpected error: {e}")
        # Try a simpler approach
        try:
            # Create a new workbook with minimal formatting
            wb = Workbook()
            ws = wb.active
            ws.title = "Extracted Text"
            
            # Just write the text line by line
            lines = text.split('\n')
            for row_idx, line in enumerate(lines, 1):
                ws.cell(row=row_idx, column=1, value=line)
            
            # Save the workbook
            wb.save(output_path)
            print(f"Excel file saved to: {output_path} (simple format)")
        except Exception as e2:
            print(f"Could not save Excel file: {e2}")

def main():
    # Parse arguments
    args = parse_arguments()
    
    # Check if file exists
    if not os.path.exists(args.file_path):
        print(f"Error: File not found: {args.file_path}")
        sys.exit(1)
    
    print(f"Processing file: {args.file_path}")
    
    # Try to install fonts in a separate thread to avoid blocking
    import threading
    font_thread = threading.Thread(target=check_and_install_fonts)
    font_thread.daemon = True  # Allow the program to exit even if thread is running
    font_thread.start()
    
    try:
        # Create AWS session with specified profile and region
        session = boto3.Session(profile_name=args.profile)
        bedrock_runtime = session.client('bedrock-runtime', region_name=args.region)
        
        # Determine the media type based on file extension
        media_type = get_media_type(args.file_path)
        print(f"Detected media type: {media_type}")
        
        # Read the file
        with open(args.file_path, 'rb') as file:
            file_content = file.read()
        
        # Create the invoke_model request body
        print(f"Calling Bedrock Runtime invoke_model API...")
        
        # For Claude models, we need to use the invoke_model API with the proper format
        if args.model.startswith('anthropic.claude'):
            # Create the request body for Claude models
            request_body = {
                "anthropic_version": "bedrock-2023-05-31",
                "max_tokens": 4000,
                "temperature": 0,
                "messages": [
                    {
                        "role": "user",
                        "content": [
                            {
                                "type": "text",
                                "text": "Please extract all the text from this document."
                            },
                            {
                                "type": "image",
                                "source": {
                                    "type": "base64",
                                    "media_type": media_type,
                                    "data": base64.b64encode(file_content).decode('utf-8')
                                }
                            }
                        ]
                    }
                ]
            }
            
            # Convert the request body to JSON
            request_body_json = json.dumps(request_body)
            
            # Make the API call to the model
            response = bedrock_runtime.invoke_model(
                modelId=args.model,
                body=request_body_json
            )
            
            # Parse the response
            response_body = json.loads(response['body'].read().decode('utf-8'))
            extracted_text = response_body['content'][0]['text']
            
        else:
            print(f"Error: Unsupported model: {args.model}")
            print("This script currently only supports Claude models.")
            sys.exit(1)
        
        # Display the extracted text
        print("\nExtracted Text:")
        print("--------------")
        print(extracted_text)
        print("\nExtraction complete")
        
        # Get base output path (without extension)
        base_output_path = os.path.splitext(args.file_path)[0]
        
        # Save in requested formats
        formats = [fmt.strip().lower() for fmt in args.formats.split(',')]
        
        # Save as text file
        if 'txt' in formats:
            output_file = f"{base_output_path}_extracted.txt"
            save_as_text(extracted_text, output_file)
        
        # Save as PDF
        if 'pdf' in formats:
            output_file = f"{base_output_path}_extracted.pdf"
            save_as_pdf(extracted_text, output_file)
        
        # Save as Word document
        if 'docx' in formats or 'doc' in formats:
            output_file = f"{base_output_path}_extracted.docx"
            save_as_docx(extracted_text, output_file)
        
        # Save as Excel file
        if 'xlsx' in formats or 'xls' in formats:
            output_file = f"{base_output_path}_extracted.xlsx"
            save_as_excel(extracted_text, output_file)
            
    except ClientError as e:
        print(f"AWS Error: {e}")
    except Exception as e:
        print(f"Unexpected error: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    main()
